#!/usr/bin/env python3
"""
Unified Impact Analysis Agent
Combines comprehensive impact analysis with tech stack recommendations
Can be used as both a service and standalone application
"""

import tempfile
import os
import json
import subprocess
import shutil
import socket
import time
import random
import urllib.request
from datetime import datetime, timedelta
from io import BytesIO
from typing import Optional, Dict, Any

# Try to import optional dependencies
try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    import docx
    from docx import Document
    DOC_SUPPORT = True
except ImportError:
    DOC_SUPPORT = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    REPORTLAB_SUPPORT = True
except ImportError:
    REPORTLAB_SUPPORT = False


class ImpactAnalysisService:
    """
    Unified Impact Analysis Service
    Provides comprehensive impact analysis and tech stack recommendations
    """
    
    def __init__(self):
        self.name = "Impact Analysis Agent"
        self.impact_categories = [
            "Business Impact",
            "Technical Impact",
            "Operational Impact",
            "Financial Impact",
            "Risk Assessment",
            "Resource Requirements",
            "Timeline Analysis",
            "Stakeholder Impact"
        ]
        
        # Configuration for Groq API (fallback)
        self.groq_api_key = os.getenv("GROQ_API_KEY", "")
        self.groq_model = os.getenv("GROQ_MODEL", "llama-3.1-8b-instant")
        self.groq_temperature = float(os.getenv("GROQ_TEMPERATURE", "0.3"))
        self.groq_url = os.getenv("GROQ_URL", "https://api.groq.com/openai/v1/chat/completions")
        
        # Try to import LLM service
        try:
            from app.services.llm import llm_service
            self.llm_service = llm_service
            self.use_internal_llm = True
        except ImportError:
            self.llm_service = None
            self.use_internal_llm = False

    # ==================== MAIN ANALYSIS METHODS ====================
    
    async def analyze_impact(
        self, 
        prd_content: str, 
        architecture_content: str, 
        github_url: str = None,
        include_tech_stack_analysis: bool = True
    ) -> dict:
        """
        Perform comprehensive impact analysis of the proposed system in ANTIGRAVITE MODE.
        
        Args:
            prd_content: Product Requirements Document content
            architecture_content: System architecture content
            github_url: Optional GitHub repository URL
            include_tech_stack_analysis: Whether to include detailed tech stack analysis
            
        Returns:
            Dictionary with report_content and file_id
        """
        try:
            from app.core.storage import register_report
            import uuid

            print("[Impact Analysis Agent] Starting ANTIGRAVITE MODE analysis...")
            
            # Extract project information
            project_info = self._extract_project_info(prd_content, architecture_content)
            
            # Clone and analyze repository if URL provided
            repo_analysis = None
            if github_url and include_tech_stack_analysis:
                print(f"[Impact Analysis Agent] Analyzing repository: {github_url}")
                repo_analysis = self.clone_and_analyze_repo(github_url)
            
            # Build comprehensive system prompt
            system_prompt = self._build_comprehensive_system_prompt()
            
            # Build user prompt with all context
            user_prompt = self._build_user_prompt(
                prd_content, 
                architecture_content, 
                github_url, 
                project_info,
                repo_analysis
            )
            
            # Generate comprehensive impact report via LLM
            report_content = await self._get_llm_response(user_prompt, system_prompt)
            
            # Store in report file
            temp_dir = tempfile.gettempdir()
            file_id = str(uuid.uuid4())
            report_filename = f"impact_analysis_{file_id}.md"
            report_path = os.path.join(temp_dir, report_filename)
            
            with open(report_path, "w", encoding="utf-8") as f:
                f.write(report_content)
            
            # Register for download
            register_report(file_id, report_path)
            
            print(f"[Impact Analysis Agent] Impact analysis completed. File ID: {file_id}")
            return {
                "report_content": report_content,
                "file_id": file_id,
                "project_info": project_info,
                "repo_analysis": repo_analysis
            }
            
        except Exception as e:
            print(f"[Impact Analysis Agent] Error: {e}")
            raise Exception(f"Failed to generate impact analysis: {e}")

    async def analyze_with_tech_stack_focus(
        self,
        repo_url: str,
        architecture_content: str,
        prd_content: str = ""
    ) -> str:
        """
        Generate tech stack focused analysis (from original impact_agent.py)
        
        Args:
            repo_url: GitHub repository URL
            architecture_content: Architecture document content
            prd_content: Optional PRD content
            
        Returns:
            Detailed tech stack analysis and recommendations
        """
        try:
            # Analyze repository
            repo_analysis = self.clone_and_analyze_repo(repo_url)
            
            # Build tech stack focused prompt
            prompt = self._build_tech_stack_prompt(
                repo_url, 
                repo_analysis, 
                architecture_content, 
                prd_content
            )
            
            # Get analysis
            return await self._get_llm_response(prompt)
            
        except Exception as e:
            return f"Error generating tech stack analysis: {str(e)}"

    # ==================== REPOSITORY ANALYSIS ====================
    
    def clone_and_analyze_repo(self, repo_url: str) -> Dict[str, Any]:
        """Clone and analyze GitHub repository with fallback options"""
        temp_dir = tempfile.mkdtemp()
        
        try:
            # Try different clone methods
            clone_commands = [
                ['git', 'clone', '--depth', '1', repo_url, temp_dir],
                ['git', 'clone', repo_url, temp_dir]
            ]
            
            clone_success = False
            last_error = None
            
            for cmd in clone_commands:
                try:
                    timeout_val = int(os.getenv("GIT_CLONE_TIMEOUT", "60"))
                    result = subprocess.run(
                        cmd, 
                        check=True, 
                        capture_output=True, 
                        text=True, 
                        timeout=timeout_val
                    )
                    clone_success = True
                    break
                except subprocess.CalledProcessError as e:
                    last_error = e
                    continue
                except subprocess.TimeoutExpired:
                    last_error = Exception("Git clone timeout")
                    continue
            
            if not clone_success:
                # Return basic analysis without cloning
                return {
                    'files': {},
                    'structure': ['Repository clone failed - analysis based on URL only'],
                    'languages': self._detect_languages_from_url(repo_url),
                    'frameworks': [],
                    'dependencies': {},
                    'clone_error': str(last_error)
                }
            
            repo_analysis = {
                'files': {},
                'structure': [],
                'languages': set(),
                'frameworks': [],
                'dependencies': {}
            }
            
            # Walk through repository
            for root, dirs, files in os.walk(temp_dir):
                # Skip hidden directories and common excludes
                dirs[:] = [d for d in dirs if not d.startswith('.') and d not in ['node_modules', '__pycache__', 'venv', 'env']]
                
                for file in files:
                    if file.startswith('.'):
                        continue
                        
                    file_path = os.path.join(root, file)
                    rel_path = os.path.relpath(file_path, temp_dir)
                    file_ext = os.path.splitext(file)[1].lower()
                    
                    repo_analysis['structure'].append(rel_path)
                    
                    # Detect programming languages
                    prog_exts = self._get_programming_extensions(repo_url)
                    if file_ext in prog_exts:
                        repo_analysis['languages'].add(file_ext)
                    
                    # Detect frameworks and dependencies
                    if file in ['package.json', 'requirements.txt', 'pom.xml', 'build.gradle', 'Cargo.toml']:
                        try:
                            with open(file_path, 'r', encoding='utf-8') as f:
                                repo_analysis['dependencies'][file] = f.read()[:1000]  # First 1000 chars
                        except:
                            pass
            
            return repo_analysis
            
        except Exception as e:
            # Return fallback analysis
            return {
                'files': {},
                'structure': ['Repository analysis unavailable'],
                'languages': self._detect_languages_from_url(repo_url),
                'frameworks': [],
                'dependencies': {},
                'analysis_error': str(e)
            }
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)

    # ==================== FILE EXTRACTION METHODS ====================
    
    def extract_text_from_file(self, file_data: bytes, filename: str) -> str:
        """Extract text from various file formats"""
        try:
            file_ext = os.path.splitext(filename)[1].lower()
            
            if file_ext == '.pdf' and PDF_SUPPORT:
                return self._extract_pdf_text(file_data)
            elif file_ext in ['.doc', '.docx'] and DOC_SUPPORT:
                return self._extract_doc_text(file_data)
            else:
                return file_data.decode('utf-8', errors='ignore')
        except Exception as e:
            raise Exception(f"Failed to extract text from {filename}: {str(e)}")
    
    def _extract_pdf_text(self, file_data: bytes) -> str:
        """Extract text from PDF files"""
        try:
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_data))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            raise Exception(f"PDF extraction failed: {str(e)}")
    
    def _extract_doc_text(self, file_data: bytes) -> str:
        """Extract text from DOC/DOCX files"""
        try:
            doc = docx.Document(BytesIO(file_data))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            raise Exception(f"DOC extraction failed: {str(e)}")

    # ==================== DOCUMENT GENERATION ====================
    
    def generate_pdf(self, content: str, title: str = "Document") -> bytes:
        """Generate PDF from text content"""
        if not REPORTLAB_SUPPORT:
            raise Exception("PDF generation not supported - missing reportlab dependency")
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Add title
        story.append(Paragraph(title, styles['Title']))
        story.append(Spacer(1, 12))
        
        # Add content paragraphs
        for line in content.split('\n'):
            if line.strip():
                story.append(Paragraph(line, styles['Normal']))
                story.append(Spacer(1, 6))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    def generate_docx(self, content: str, title: str = "Document") -> bytes:
        """Generate DOCX from text content"""
        if not DOC_SUPPORT:
            raise Exception("DOCX generation not supported - missing python-docx dependency")
        
        doc = Document()
        doc.add_heading(title, 0)
        
        for line in content.split('\n'):
            if line.strip():
                doc.add_paragraph(line)
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    # ==================== PROMPT BUILDING ====================
    
    def _build_comprehensive_system_prompt(self) -> str:
        """Build comprehensive system prompt for impact analysis"""
        return """You are operating in **ANTIGRAVITE MODE**.

You are **Agent-3: Principal Impact Analyst & Enterprise Architecture Evaluator**.

Your responsibility is to generate a **professional Impact Analysis Report**
and produce a **PDF-ready output** based on real repository, architecture,
and PRD analysis.

==================================================
AGENT-3 OBJECTIVE
==================================================
Generate a **comprehensive Impact Analysis Report** that evaluates:
• Technical impact
• Business impact
• Scalability impact
• Security impact
• Cost & operational impact
• Technology risk & alternatives

The final output must be **structured, executive-ready, and PDF-exportable**.

==================================================
MANDATORY SECTIONS
==================================================

1️⃣ PROJECT SUMMARY  
• High-level overview of the system
• Core business goals
• Supported user flows
• Platform scope (web/mobile/admin)

2️⃣ ARCHITECTURE OVERVIEW  
• Text-based architecture diagram
• User → Frontend → Backend → Database → Integrations
• Clear data flow explanation

3️⃣ TECH STACK IMPACT ANALYSIS  
For EACH layer (Frontend, Backend, Database):
• Why it fits the requirements
• Strengths
• Limitations
• Long-term impact

4️⃣ TECHNOLOGY INTEGRATION IMPACT  
• Frontend ↔ Backend communication
• Backend ↔ Database interactions
• External services (payments, email, SMS)
• Performance and reliability implications

5️⃣ ALTERNATIVE TECHNOLOGY STACKS  
Evaluate alternatives with:
• Pros
• Cons
• Migration impact
• Cost & learning curve

Sections:
• Backend alternatives
• Database alternatives
• Frontend alternatives

6️⃣ INTELLIGENT DATABASE IMPACT  
• Schema scalability
• Indexing effectiveness
• Transaction safety
• Data growth readiness
• Operational complexity

7️⃣ API DESIGN IMPACT  
• REST compliance
• Endpoint scalability
• Validation & security exposure
• Versioning readiness
• Maintenance impact

8️⃣ SECURITY IMPACT ANALYSIS  
• Authentication model
• Authorization control
• Data protection
• Threat surface
• Compliance readiness

9️⃣ PERFORMANCE & SCALABILITY IMPACT  
• Expected traffic handling
• Horizontal/vertical scaling
• Caching opportunities
• Bottleneck risks

🔟 BUSINESS & COST IMPACT  
• Development effort
• Operational cost
• Maintenance overhead
• Vendor lock-in risk
• ROI alignment

1️⃣1️⃣ RISK ANALYSIS  
Create a table:
• Risk
• Severity
• Probability
• Mitigation strategy

1️⃣2️⃣ IMPLEMENTATION ROADMAP IMPACT  
Analyze:
• Phase-wise risk
• Technical dependencies
• Failure points
• Optimization opportunities

1️⃣3️⃣ FINAL IMPACT SCORECARD  
Score (1–10):
• Technical robustness
• Scalability
• Security
• Maintainability
• Business alignment

1️⃣4️⃣ RECOMMENDATIONS & NEXT STEPS  
• Architecture improvements
• Technology upgrades
• Risk mitigation actions
• Optimization roadmap

==================================================
OUTPUT & PDF RULES
==================================================
• Formal enterprise language
• Clear headings
• Tables where applicable
• No emojis in final report
• No casual tone
• No assumptions without analysis
• Ready for direct PDF export
"""

    def _build_user_prompt(
        self,
        prd_content: str,
        architecture_content: str,
        github_url: str,
        project_info: dict,
        repo_analysis: Optional[dict] = None
    ) -> str:
        """Build user prompt with all context"""
        
        repo_context = ""
        if repo_analysis:
            languages = list(repo_analysis.get('languages', []))[:5]
            structure = repo_analysis.get('structure', [])[:10]
            repo_context = f"""
REPOSITORY ANALYSIS:
Languages Detected: {', '.join(languages)}
File Structure Sample: {', '.join(structure)}
Dependencies: {json.dumps(repo_analysis.get('dependencies', {}), indent=2)[:500]}
"""
        
        return f"""
ANALYZE THE FOLLOWING CONTEXT AND GENERATE THE IMPACT REPORT:

PRD CONTENT:
{prd_content[:2000] if prd_content else 'Not provided'}

ARCHITECTURE BLUEPRINT (Agent-2 Output):
{architecture_content[:3000] if architecture_content else 'Not provided'}

{repo_context}

GITHUB REPOSITORY: {github_url or 'N/A'}
CURRENT TIMESTAMP: {datetime.utcnow().isoformat()}
PROJECT SCALE: {project_info['scale']}
COMPLEXITY: {project_info['complexity']}
FEATURES COUNT: {len(project_info['features'])}
INTEGRATIONS: {', '.join(project_info['integrations'])}
COMPLIANCE: {', '.join(project_info['compliance_requirements'])}
"""

    def _build_tech_stack_prompt(
        self,
        repo_url: str,
        repo_analysis: dict,
        architecture_content: str,
        prd_content: str
    ) -> str:
        """Build tech stack focused prompt"""
        
        # Truncate architecture content if too long
        max_arch_chars = int(os.getenv("MAX_ARCH_CHARS", "3000"))
        if len(architecture_content) > max_arch_chars:
            architecture_content = architecture_content[:max_arch_chars] + "\n[Content truncated due to length]"
        
        # Extract key repo info
        frontend_exts = self._get_frontend_extensions(architecture_content)
        frontend_files = [
            f for f in repo_analysis['structure'][:5] 
            if any(f.endswith(ext.strip()) for ext in frontend_exts)
        ]
        languages = list(repo_analysis['languages'])[:3]
        
        return f"""Analyze this frontend repo and architecture document to provide comprehensive tech stack justification and alternatives:

REPO: {repo_url}
LANGUAGES: {', '.join(languages)}
FRONTEND FILES: {', '.join(frontend_files)}

ARCHITECTURE:
{architecture_content}

PRD CONTENT:
{prd_content or 'Not provided'}

Generate output with:

# PROJECT SUMMARY
**Repository**: {repo_url}
[Generate a concise project summary based on the repository analysis and architecture document]

# ARCHITECTURE DIAGRAM
[Create an ASCII-based system architecture diagram showing the flow between frontend, backend, database, and external services. Use boxes, arrows, and clear labels]

# TECH STACK JUSTIFICATION
[For each technology mentioned in the architecture document, provide detailed justification why it's suitable for this project, including pros/cons and fit with requirements]

# ALTERNATIVE TECH STACKS
## Backend Alternatives:
- Option 1: [Framework/Language] - Pros, Cons, Performance, Scalability
- Option 2: [Framework/Language] - Pros, Cons, Performance, Scalability
- Option 3: [Framework/Language] - Pros, Cons, Performance, Scalability

## Database Alternatives:
- Option 1: [Database Type] - Use cases, Performance, Cost
- Option 2: [Database Type] - Use cases, Performance, Cost
- Option 3: [Database Type] - Use cases, Performance, Cost

# DATABASE SCHEMA DESIGN
[Based on architecture requirements, design database tables with fields, data types, relationships, and indexes]

# RECOMMENDED API ENDPOINTS
[Based on the architecture requirements, suggest complete API specification with:
- HTTP methods (GET, POST, PUT, DELETE)
- Full endpoint paths
- Request/response body formats
- Authentication requirements
- Input validation rules
- Error response formats]

## Input Fields for Each Endpoint:
[For each API endpoint, specify:
- Required input fields with data types
- Optional input fields with defaults
- Validation rules (min/max length, format, etc.)
- Example request payloads
- Field descriptions and purposes]

# DETAILED PROJECT CONSTRUCTION GUIDE
[Use the BEST ALTERNATIVE tech stack from the alternatives analysis above. If architecture document tech stack is clearly superior, use that instead. Clearly state which tech stack combination you're using for this guide.]

## Phase 1: Environment Setup
1. Install [SPECIFIC BACKEND FRAMEWORK] development environment
2. Setup [SPECIFIC DATABASE] server and tools
3. Initialize Git repository with proper .gitignore
4. Create project folder structure for chosen tech stack
5. Setup package managers and dependency files
6. Configure environment variables template

## Phase 2: Backend Development
1. Install [SPECIFIC FRAMEWORK] and create initial project structure
2. Configure [SPECIFIC DATABASE] connection with credentials
3. Create database models/entities based on schema design above
4. Create input validation schemas for all API endpoints
5. Implement each API endpoint with detailed input/output handling
6. Setup JWT/OAuth authentication with input validation
7. Add middleware for logging, CORS, rate limiting, input sanitization
8. Create comprehensive test suites for all endpoints and input validation

## Phase 3: Database Implementation
1. Setup [SPECIFIC DATABASE TYPE] server (local/cloud)
2. Run database migrations to create schema from design above
3. Create seed scripts with sample data
4. Add database indexes for performance optimization
5. Setup automated backup procedures
6. Configure connection pooling and optimization

## Phase 4: Integration
1. Connect existing frontend to new backend APIs
2. Update frontend API calls to match new endpoints
3. Implement error handling and loading states
4. Add authentication flow integration
5. Test all frontend-backend data flows
6. Optimize API response times and caching

Provide SPECIFIC commands, code snippets, and configuration examples for the CHOSEN tech stack in each step."""

    # ==================== PROJECT INFO EXTRACTION ====================
    
    def _extract_project_info(self, prd_content: str, architecture_content: str = "") -> dict:
        """Extract key project information from project documentation."""
        
        # Extract project characteristics
        project_info = {
            "complexity": "Medium",
            "scale": "Medium",
            "user_base": "Unknown",
            "features": [],
            "integrations": [],
            "compliance_requirements": []
        }
        
        combined_content = (prd_content or "") + "\n" + (architecture_content or "")
        content_lower = combined_content.lower()
        
        # Analyze complexity indicators
        complexity_indicators = {
            "High": ["machine learning", "ai", "real-time", "microservices", "blockchain", "streaming", "security-critical"],
            "Medium": ["api", "database", "authentication", "payment", "notification", "dashboard"],
            "Low": ["crud", "simple", "basic", "static", "profile"]
        }
        
        for level, indicators in complexity_indicators.items():
            if any(indicator in content_lower for indicator in indicators):
                project_info["complexity"] = level
                break
        
        # Extract features
        lines = combined_content.split('\n')
        for line in lines:
            if any(keyword in line.lower() for keyword in ['feature', 'functionality', 'requirement', 'endpoint', 'module']):
                clean_feature = line.strip().lstrip('-•*# ').strip()
                if 8 < len(clean_feature) < 150:
                    project_info["features"].append(clean_feature)
        
        # Ensure we have some features to show
        if not project_info["features"]:
            project_info["features"] = ["Core API Development", "User Management System", "Data Persistence Layer"]
        
        # Estimate scale based on features
        feature_count = len(project_info["features"])
        if feature_count > 12:
            project_info["scale"] = "Large"
        elif feature_count > 5:
            project_info["scale"] = "Medium"
        else:
            project_info["scale"] = "Small"
        
        # Detect integrations
        integration_keywords = ["payment", "email", "sms", "social", "api", "third-party", "gateway"]
        for keyword in integration_keywords:
            if keyword in content_lower:
                project_info["integrations"].append(keyword.title())
        
        # Detect compliance requirements
        compliance_keywords = ["gdpr", "hipaa", "pci", "sox", "privacy", "security", "encryption"]
        for keyword in compliance_keywords:
            if keyword in content_lower:
                project_info["compliance_requirements"].append(keyword.upper())
        
        return project_info

    # ==================== HELPER METHODS ====================
    
    def _detect_languages_from_url(self, repo_url: str) -> set:
        """Detect likely programming languages from repository URL without LLM"""
        return set(['.js', '.html', '.css', '.jsx', '.ts', '.tsx'])
    
    def _get_frontend_extensions(self, architecture_content: str) -> list:
        """Get frontend file extensions without LLM"""
        return ['.js', '.jsx', '.ts', '.tsx', '.vue', '.html', '.css']
    
    def _get_programming_extensions(self, repo_url: str) -> list:
        """Get programming extensions without LLM"""
        return ['.py', '.js', '.jsx', '.ts', '.tsx', '.java', '.go', '.rs', '.php', '.rb', '.cs', '.html', '.css']
    
    def _estimate_development_time(self, project_info: dict) -> str:
        """Estimate development time based on project characteristics."""
        
        base_time = {
            "Small": 3,
            "Medium": 6,
            "Large": 12
        }
        
        complexity_adjustment = {
            "Low": 0.8,
            "Medium": 1.0,
            "High": 1.5
        }
        
        scale = project_info["scale"]
        complexity = project_info["complexity"]
        
        estimated_months = base_time.get(scale, 6) * complexity_adjustment.get(complexity, 1.0)
        
        return f"{int(estimated_months)}-{int(estimated_months * 1.2)} months"

    # ==================== LLM INTEGRATION ====================
    
    async def _get_llm_response(self, user_prompt: str, system_prompt: str = "") -> str:
        """Get response from LLM (internal service or Groq API)"""
        
        # Try internal LLM service first
        if self.use_internal_llm and self.llm_service:
            try:
                return await self.llm_service.get_response(user_prompt, system_prompt)
            except Exception as e:
                print(f"[Impact Analysis] Internal LLM failed: {e}, falling back to Groq")
        
        # Fallback to Groq API
        if self.groq_api_key:
            try:
                return self._call_groq_api(user_prompt, system_prompt)
            except Exception as e:
                print(f"[Impact Analysis] Groq API failed: {e}")
                raise
        
        raise Exception("No LLM service available. Configure either internal LLM or GROQ_API_KEY")
    
    def _call_groq_api(self, user_prompt: str, system_prompt: str = "", max_retries: int = 3) -> str:
        """Make API call to Groq with retry logic for rate limits"""
        if not self.groq_api_key:
            raise Exception("GROQ_API_KEY not configured")
        
        messages = []
        if system_prompt:
            messages.append({"role": "system", "content": system_prompt})
        messages.append({"role": "user", "content": user_prompt})
        
        for attempt in range(max_retries):
            try:
                data = {
                    "messages": messages,
                    "model": self.groq_model,
                    "temperature": self.groq_temperature,
                    "max_tokens": int(os.getenv("GROQ_MAX_TOKENS", "4500"))
                }
                
                headers = {
                    'Authorization': f'Bearer {self.groq_api_key}',
                    'Content-Type': 'application/json'
                }
                
                req = urllib.request.Request(
                    self.groq_url,
                    data=json.dumps(data).encode('utf-8'),
                    headers=headers
                )
                
                with urllib.request.urlopen(req) as response:
                    result = json.loads(response.read().decode('utf-8'))
                    return result['choices'][0]['message']['content']
                    
            except urllib.error.HTTPError as e:
                error_body = e.read().decode('utf-8')
                
                # Handle rate limit specifically
                if e.code == 429:
                    if attempt < max_retries - 1:
                        wait_time = 10 + (attempt * 15) + random.uniform(0, 5)
                        print(f"[Impact Analysis] Rate limited, waiting {wait_time}s...")
                        time.sleep(wait_time)
                        continue
                    else:
                        raise Exception(f"Rate limit exceeded after {max_retries} attempts. Please wait 60 seconds and try again.")
                else:
                    raise Exception(f"Groq API Error {e.code}: {error_body}")
            except Exception as e:
                if attempt < max_retries - 1:
                    time.sleep(1)
                    continue
                raise Exception(f"API call failed: {str(e)}")
        
        raise Exception("Max retries exceeded")


# Create singleton instance
impact_analysis_service = ImpactAnalysisService()


# ==================== STANDALONE SERVER (OPTIONAL) ====================

def create_standalone_app():
    """
    Create standalone FastAPI application for Impact Analysis Agent
    This allows the service to run independently if needed
    """
    try:
        from fastapi import FastAPI, File, UploadFile, HTTPException
        from fastapi.responses import HTMLResponse, StreamingResponse
        from fastapi.middleware.cors import CORSMiddleware
        from pydantic import BaseModel
    except ImportError:
        print("FastAPI not available. Standalone mode disabled.")
        return None
    
    app = FastAPI(
        title="Impact Analysis Agent", 
        description="Comprehensive Impact Analysis and Tech Stack Recommendations",
        docs_url="/api/docs", 
        redoc_url="/api/redoc"
    )
    
    # Add CORS middleware
    app.add_middleware(
        CORSMiddleware,
        allow_origins=["*"],
        allow_credentials=True,
        allow_methods=["*"],
        allow_headers=["*"],
    )
    
    # Pydantic models
    class AnalysisRequest(BaseModel):
        repo_url: str
        architecture_content: str
        prd_content: str = ""
        include_tech_stack: bool = True
    
    class FileUploadResponse(BaseModel):
        success: bool
        extracted_text: str = ""
        filename: str = ""
        error: str = ""
    
    class AnalysisResponse(BaseModel):
        success: bool
        analysis: str = ""
        document_id: str = ""
        timestamp: str = ""
        error: str = ""
    
    # Routes
    @app.get("/", response_class=HTMLResponse)
    async def read_root():
        """Serve the main HTML page"""
        html_content = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Impact Analysis Agent</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 40px; background: #f5f5f5; }
        .container { max-width: 800px; margin: 0 auto; background: white; padding: 30px; border-radius: 10px; box-shadow: 0 4px 6px rgba(0,0,0,0.1); }
        h1 { color: #333; text-align: center; margin-bottom: 30px; }
        .form-group { margin-bottom: 20px; }
        label { display: block; margin-bottom: 5px; font-weight: bold; color: #555; }
        input[type="url"], input[type="file"] { width: 100%; padding: 10px; border: 2px solid #ddd; border-radius: 5px; }
        button { background: #007bff; color: white; padding: 12px 30px; border: none; border-radius: 5px; cursor: pointer; font-size: 16px; }
        button:hover { background: #0056b3; }
        .status { margin-top: 20px; padding: 10px; border-radius: 5px; }
        .success { background: #d4edda; color: #155724; }
        .error { background: #f8d7da; color: #721c24; }
    </style>
</head>
<body>
    <div class="container">
        <h1>🚀 Impact Analysis Agent</h1>
        <p style="text-align: center; color: #666;">Comprehensive Impact Analysis & Tech Stack Recommendations</p>
        <form id="analysisForm">
            <div class="form-group">
                <label>GitHub Repository URL:</label>
                <input type="url" id="repoUrl" placeholder="https://github.com/username/repository" required>
            </div>
            <div class="form-group">
                <label>Architecture Document:</label>
                <input type="file" id="archFile" accept=".pdf,.doc,.docx,.txt,.json" required>
            </div>
            <div class="form-group">
                <label>PRD Document (Optional):</label>
                <input type="file" id="prdFile" accept=".pdf,.doc,.docx,.txt,.json">
            </div>
            <button type="submit">🔍 Analyze Project</button>
        </form>
        <div id="status"></div>
    </div>
</body>
</html>
        """
        return HTMLResponse(content=html_content)
    
    @app.post("/upload-file", response_model=FileUploadResponse)
    async def upload_file(file: UploadFile = File(...)):
        """Handle file upload and text extraction"""
        try:
            file_data = await file.read()
            extracted_text = impact_analysis_service.extract_text_from_file(file_data, file.filename)
            
            return FileUploadResponse(
                success=True,
                extracted_text=extracted_text,
                filename=file.filename
            )
        except Exception as e:
            return FileUploadResponse(
                success=False,
                error=str(e)
            )
    
    @app.post("/analyze", response_model=AnalysisResponse)
    async def analyze_project(request: AnalysisRequest):
        """Handle analysis requests"""
        try:
            result = await impact_analysis_service.analyze_impact(
                prd_content=request.prd_content,
                architecture_content=request.architecture_content,
                github_url=request.repo_url,
                include_tech_stack_analysis=request.include_tech_stack
            )
            
            return AnalysisResponse(
                success=True,
                analysis=result["report_content"],
                document_id=result["file_id"],
                timestamp=datetime.now().isoformat()
            )
            
        except Exception as e:
            return AnalysisResponse(
                success=False,
                error=str(e)
            )
    
    return app


# Entry point for standalone mode
if __name__ == "__main__":
    import uvicorn
    
    app = create_standalone_app()
    if app:
        def find_available_port(start_port=8090, max_attempts=10):
            for port in range(start_port, start_port + max_attempts):
                try:
                    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
                        s.bind(('127.0.0.1', port))
                        return port
                except OSError:
                    continue
            return 8090
        
        port = find_available_port()
        print(f"Starting Impact Analysis Agent on port {port}")
        uvicorn.run(app, host="0.0.0.0", port=port)
    else:
        print("Cannot start standalone server - FastAPI not installed")